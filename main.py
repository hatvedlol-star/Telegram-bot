
import os
import time
import tempfile
import requests

import numpy as np
from PIL import Image, ImageEnhance, ImageFilter
import pytesseract
import speech_recognition as sr
from telegram import InlineKeyboardButton, InlineKeyboardMarkup, ReplyKeyboardMarkup, ReplyKeyboardRemove
from telegram.ext import Updater, CommandHandler, MessageHandler, Filters, CallbackQueryHandler
import subprocess
from docx import Document
from docx.shared import Inches
import fitz  # PyMuPDF

TOKEN = "8207135505:AAH0eOZjnCzFVIKN5ilWSAxEIWfiu2wPBDM"
CHANNEL_USERNAME = "@News_SirMob"
CLOUDCONVERT_API_KEY = "eyJ0eXAiOiJKV1QiLCJhbGciOiJSUzI1NiJ9.eyJhdWQiOiIxIiwianRpIjoiYzkyZDlkNTE4YzY3ZmMxMTkyZGNiYjk5N2JhN2IxZmZhNWE1ZDdjODA0Y2RiZDRlMzdhNTlkYTI5MmQ3ZmViYThhODE3MTQ1NmUxMmNkYWYiLCJpYXQiOjE3NTYyNDI4NjkuODc5NDc0LCJuYmYiOjE3NTYyNDI4NjkuODc5NDc2LCJleHAiOjQ5MTE5MTY0NjkuODczNTE5LCJzdWIiOiI3Mjc1NTMzMCIsInNjb3BlcyI6WyJ1c2VyLnJlYWQiLCJ1c2VyLndyaXRlIiwidGFzay5yZWFkIiwidGFzay53cml0ZSIsIndlYmhvb2sucmVhZCIsIndlYmhvb2sud3JpdGUiLCJwcmVzZXQucmVhZCIsInByZXNldC53cml0ZSJdfQ.fEXNM-nXsL2ROOnjICMf0UzAAye6DFivfLZT32Xg8Now-EdqjGGVRN5uy4S54A1aUnng6uYDVASlqxVcXIYcFuRjbnI7adrsEQuOfsUykvAUxbv4l5vG296dMQvx9ojwXW0QL4aq2RVrFGMoQVnps0jJDUM88xhewO1ifrA1vbSSAY-7z7WXnIfCbYx4gZQKn21wVC8YB9W3ZSWf5yEv7gX7n7H61cGzFlPs138O5M0K8TCCcBXdf8y1x3ptGaa1Un5pfP08eVe0ajdeCpaCqQUuYk1QKqcR-EuVsX2GZ7BcV18fQBtAkPL2W6-T4KkmVN3lsXh92ox1H6J427hBg8kmCgIf-igx4zhjyWyWC815ASV1wBGhOQ6bwggGe6c29LnSQPf0o-572yiUuijvFfzoef-u14fyPsTbAwFyJ7nMtCYOL7B0TJrMKZv901jYxPtDnhJTI3hLVC1Qq936AxuH0CxgehW547lH7byT00ROz5f9JNWMHNts5ALVC9sLnr8M0MyhtaOUsm7Ol-CEanDzBaJPG6T5pTHNZH5MWcPNJ8QM7N0nc2RexQw113F3TROtAcHJAOpDCGl9SD65aJ2TutiOOi2Kj5UIr9jn7iJaKrOmh95IMmaiWi0F3HnhtWJ7F7x9ftm3uaFxjOtahgBbv10"

# Dictionary to track user operations
user_operations = {}

def is_subscribed(bot, user_id):
    try:
        member = bot.get_chat_member(CHANNEL_USERNAME, user_id)
        return member.status in ["member", "administrator", "creator"]
    except:
        return False

def show_subscription_message(update):
    keyboard = [[InlineKeyboardButton("📢 اشترك في قناة المطور", url=f"https://t.me/{CHANNEL_USERNAME.lstrip('@')}")]]
    update.message.reply_text("اشترك في القناة أولاً.", reply_markup=InlineKeyboardMarkup(keyboard))

def build_main_keyboard():
    keyboard = [
        ["🎬 فيديو ⬅️ صوت", "🖼️ صور ⬅️ PDF"],
        ["📄 PDF ⬅️ Word", "📝 Word ⬅️ PDF"],
        ["🎤 صوت ⬅️ نص"]
    ]
    return ReplyKeyboardMarkup(keyboard, resize_keyboard=True)

def build_cancel_keyboard():
    keyboard = [["🛑 إيقاف العملية"]]
    return ReplyKeyboardMarkup(keyboard, resize_keyboard=True)

def is_user_busy(user_id):
    return user_id in user_operations

def set_user_operation(user_id, operation):
    user_operations[user_id] = operation

def clear_user_operation(user_id):
    if user_id in user_operations:
        del user_operations[user_id]

def start(update, context):
    user_id = update.effective_user.id
    
    if not is_subscribed(context.bot, user_id):
        show_subscription_message(update)
        return
    
    clear_user_operation(user_id)
    update.message.reply_text(
        "مرحباً! اختر العملية التي تريد تنفيذها:",
        reply_markup=build_main_keyboard()
    )

def handle_message(update, context):
    user_id = update.effective_user.id
    text = update.message.text
    
    if not is_subscribed(context.bot, user_id):
        show_subscription_message(update)
        return
    
    if text == "🛑 إيقاف العملية":
        clear_user_operation(user_id)
        update.message.reply_text("تم إيقاف العملية.", reply_markup=build_main_keyboard())
        return
    
    # Handle main menu options
    if text == "🎬 فيديو ⬅️ صوت":
        set_user_operation(user_id, "video_to_audio")
        update.message.reply_text("أرسل الفيديو الذي تريد تحويله إلى صوت.", reply_markup=build_cancel_keyboard())
    elif text == "🖼️ صور ⬅️ PDF":
        set_user_operation(user_id, "images_to_pdf")
        update.message.reply_text("أرسل الصور التي تريد تحويلها إلى PDF.", reply_markup=build_cancel_keyboard())
    elif text == "📄 PDF ⬅️ Word":
        set_user_operation(user_id, "pdf_to_word")
        update.message.reply_text("أرسل ملف PDF الذي تريد تحويله إلى Word.", reply_markup=build_cancel_keyboard())
    elif text == "📝 Word ⬅️ PDF":
        set_user_operation(user_id, "word_to_pdf")
        update.message.reply_text("أرسل ملف Word الذي تريد تحويله إلى PDF.", reply_markup=build_cancel_keyboard())
    elif text == "🎤 صوت ⬅️ نص":
        set_user_operation(user_id, "audio_to_text")
        update.message.reply_text("أرسل الملف الصوتي الذي تريد تحويله إلى نص.", reply_markup=build_cancel_keyboard())

def handle_video(update, context):
    user_id = update.effective_user.id
    
    if not is_subscribed(context.bot, user_id):
        show_subscription_message(update)
        return
    
    if not is_user_busy(user_id) or user_operations[user_id] != "video_to_audio":
        update.message.reply_text("اختر 'فيديو ⬅️ صوت' من القائمة أولاً.")
        return
    
    try:
        update.message.reply_text("جاري تحويل الفيديو إلى صوت...")
        
        video_file = update.message.video.get_file()
        with tempfile.NamedTemporaryFile(suffix='.mp4', delete=False) as temp_video:
            video_file.download(temp_video.name)
            
            audio_path = temp_video.name.replace('.mp4', '.mp3')
            
            # Convert video to audio using ffmpeg
            subprocess.run(['ffmpeg', '-i', temp_video.name, '-q:a', '0', '-map', 'a', audio_path], check=True)
            
            with open(audio_path, 'rb') as audio_file:
                update.message.reply_audio(audio_file, caption="تم تحويل الفيديو إلى صوت بنجاح!")
            
            os.unlink(temp_video.name)
            os.unlink(audio_path)
            
        clear_user_operation(user_id)
        update.message.reply_text("اختر عملية أخرى:", reply_markup=build_main_keyboard())
        
    except Exception as e:
        update.message.reply_text(f"حدث خطأ: {str(e)}")
        clear_user_operation(user_id)

def handle_photo(update, context):
    user_id = update.effective_user.id
    
    if not is_subscribed(context.bot, user_id):
        show_subscription_message(update)
        return
    
    if not is_user_busy(user_id) or user_operations[user_id] != "images_to_pdf":
        update.message.reply_text("اختر 'صور ⬅️ PDF' من القائمة أولاً.")
        return
    
    try:
        update.message.reply_text("جاري معالجة الصورة...")
        
        # Initialize images list if not exists
        if f"{user_id}_images" not in user_operations:
            user_operations[f"{user_id}_images"] = []
        
        photo_file = update.message.photo[-1].get_file()
        with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as temp_photo:
            photo_file.download(temp_photo.name)
            user_operations[f"{user_id}_images"].append(temp_photo.name)
        
        keyboard = [
            ["إضافة صورة أخرى", "إنهاء وتحويل إلى PDF"],
            ["🛑 إيقاف العملية"]
        ]
        reply_markup = ReplyKeyboardMarkup(keyboard, resize_keyboard=True)
        
        update.message.reply_text(
            f"تم إضافة الصورة ({len(user_operations[f'{user_id}_images'])} صور حتى الآن).\nماذا تريد أن تفعل؟",
            reply_markup=reply_markup
        )
        
    except Exception as e:
        update.message.reply_text(f"حدث خطأ: {str(e)}")

def handle_document(update, context):
    user_id = update.effective_user.id
    
    if not is_subscribed(context.bot, user_id):
        show_subscription_message(update)
        return
    
    if not is_user_busy(user_id):
        update.message.reply_text("اختر عملية من القائمة أولاً.")
        return
    
    operation = user_operations[user_id]
    
    try:
        if operation == "pdf_to_word":
            update.message.reply_text("جاري تحويل PDF إلى Word...")
            
            doc_file = update.message.document.get_file()
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
                doc_file.download(temp_pdf.name)
                
                # Extract text from PDF
                pdf_document = fitz.open(temp_pdf.name)
                text = ""
                for page_num in range(pdf_document.page_count):
                    page = pdf_document[page_num]
                    text += page.get_text()
                
                # Create Word document
                doc = Document()
                doc.add_paragraph(text)
                
                word_path = temp_pdf.name.replace('.pdf', '.docx')
                doc.save(word_path)
                
                with open(word_path, 'rb') as word_file:
                    update.message.reply_document(word_file, caption="تم تحويل PDF إلى Word بنجاح!")
                
                os.unlink(temp_pdf.name)
                os.unlink(word_path)
                
        elif operation == "word_to_pdf":
            update.message.reply_text("جاري تحويل Word إلى PDF...")
            
            doc_file = update.message.document.get_file()
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_doc:
                doc_file.download(temp_doc.name)
                
                # Read Word document
                doc = Document(temp_doc.name)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                
                # Create PDF using reportlab
                from reportlab.pdfgen import canvas
                from reportlab.lib.pagesizes import letter
                
                pdf_path = temp_doc.name.replace('.docx', '.pdf')
                c = canvas.Canvas(pdf_path, pagesize=letter)
                width, height = letter
                
                y = height - 50
                for line in text.split('\n'):
                    if y < 50:
                        c.showPage()
                        y = height - 50
                    c.drawString(50, y, line[:80])  # Limit line length
                    y -= 20
                
                c.save()
                
                with open(pdf_path, 'rb') as pdf_file:
                    update.message.reply_document(pdf_file, caption="تم تحويل Word إلى PDF بنجاح!")
                
                os.unlink(temp_doc.name)
                os.unlink(pdf_path)
        
        clear_user_operation(user_id)
        update.message.reply_text("اختر عملية أخرى:", reply_markup=build_main_keyboard())
        
    except Exception as e:
        update.message.reply_text(f"حدث خطأ: {str(e)}")
        clear_user_operation(user_id)

def handle_audio(update, context):
    user_id = update.effective_user.id
    
    if not is_subscribed(context.bot, user_id):
        show_subscription_message(update)
        return
    
    if not is_user_busy(user_id) or user_operations[user_id] != "audio_to_text":
        update.message.reply_text("اختر 'صوت ⬅️ نص' من القائمة أولاً.")
        return
    
    try:
        update.message.reply_text("جاري تحويل الصوت إلى نص...")
        
        audio_file = None
        original_suffix = '.ogg'
        
        if update.message.audio:
            audio_file = update.message.audio.get_file()
            original_suffix = '.mp3'
        elif update.message.voice:
            audio_file = update.message.voice.get_file()
            original_suffix = '.ogg'
        else:
            update.message.reply_text("نوع الملف غير مدعوم. يرجى إرسال ملف صوتي.")
            return
        
        with tempfile.NamedTemporaryFile(suffix=original_suffix, delete=False) as temp_audio:
            audio_file.download(temp_audio.name)
            
            # Convert audio to wav format for better compatibility
            wav_path = temp_audio.name.replace(original_suffix, '.wav')
            subprocess.run(['ffmpeg', '-i', temp_audio.name, '-ar', '16000', '-ac', '1', wav_path], check=True)
            
            # Use speech recognition
            recognizer = sr.Recognizer()
            
            with sr.AudioFile(wav_path) as source:
                audio_data = recognizer.record(source)
                
                # Try Arabic first
                try:
                    text = recognizer.recognize_google(audio_data, language='ar-SA')
                except sr.UnknownValueError:
                    try:
                        # Try English if Arabic fails
                        text = recognizer.recognize_google(audio_data, language='en-US')
                    except sr.UnknownValueError:
                        text = ""
                except sr.RequestError:
                    text = ""
            
            # Clean up wav file
            if os.path.exists(wav_path):
                os.unlink(wav_path)
            
            if text:
                update.message.reply_text(f"النص المستخرج:\n\n{text}")
            else:
                update.message.reply_text("لم يتم التعرف على أي كلام واضح في الملف الصوتي. تأكد من وضوح الصوت وجودة التسجيل.")
            
            # Clean up files
            os.unlink(temp_audio.name)
        
        clear_user_operation(user_id)
        update.message.reply_text("اختر عملية أخرى:", reply_markup=build_main_keyboard())
        
    except Exception as e:
        update.message.reply_text(f"حدث خطأ في تحويل الصوت: {str(e)}")
        clear_user_operation(user_id)

def handle_special_commands(update, context):
    user_id = update.effective_user.id
    text = update.message.text
    
    if text == "إضافة صورة أخرى":
        update.message.reply_text("أرسل الصورة التالية.")
    elif text == "إنهاء وتحويل إلى PDF":
        if f"{user_id}_images" in user_operations and user_operations[f"{user_id}_images"]:
            try:
                update.message.reply_text("جاري إنشاء ملف PDF...")
                
                images = user_operations[f"{user_id}_images"]
                pdf_path = f"/tmp/images_{user_id}.pdf"
                
                # Convert images to PDF
                image_list = []
                for img_path in images:
                    img = Image.open(img_path)
                    if img.mode != 'RGB':
                        img = img.convert('RGB')
                    image_list.append(img)
                
                image_list[0].save(pdf_path, save_all=True, append_images=image_list[1:])
                
                with open(pdf_path, 'rb') as pdf_file:
                    update.message.reply_document(pdf_file, caption="تم تحويل الصور إلى PDF بنجاح!")
                
                # Clean up
                for img_path in images:
                    os.unlink(img_path)
                os.unlink(pdf_path)
                del user_operations[f"{user_id}_images"]
                
                clear_user_operation(user_id)
                update.message.reply_text("اختر عملية أخرى:", reply_markup=build_main_keyboard())
                
            except Exception as e:
                update.message.reply_text(f"حدث خطأ: {str(e)}")
        else:
            update.message.reply_text("لم يتم العثور على صور.")
    else:
        handle_message(update, context)

def main():
    updater = Updater(TOKEN, use_context=True)
    dp = updater.dispatcher
    
    dp.add_handler(CommandHandler("start", start))
    dp.add_handler(MessageHandler(Filters.video, handle_video))
    dp.add_handler(MessageHandler(Filters.photo, handle_photo))
    dp.add_handler(MessageHandler(Filters.document, handle_document))
    dp.add_handler(MessageHandler(Filters.audio | Filters.voice, handle_audio))
    dp.add_handler(MessageHandler(Filters.text & ~Filters.command, handle_special_commands))
    
    print("البوت يعمل الآن...")
    updater.start_polling()
    updater.idle()

if __name__ == "__main__":
    main()
